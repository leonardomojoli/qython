# information-department-adapter/utils/file_handler

import os
import subprocess
import tempfile
from PIL import Image
import pytesseract
import glob  # Adicionado para listar arquivos de imagem gerados
import logging # Adicionado logging
from .transcription_service import transcribe_audio # Importar transcriber
from ...config import Config # Importar Config
import fitz

# Define o caminho para o executável do Tesseract para evitar problemas de PATH
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# Configurar logging
logger = logging.getLogger(__name__)

os.environ["TESSDATA_PREFIX"] = "/usr/share/tesseract-ocr/5/tessdata/"

def limpar_arquivos_locais(arquivos):
    """
    Remove arquivos locais especificados.
    """
    for arquivo in arquivos:
        if os.path.exists(arquivo):
            os.remove(arquivo)
            logger.debug(f"Arquivo local {arquivo} removido.")
        else:
            logger.debug(f"Arquivo {arquivo} não encontrado para remoção.")

def converter_para_imagem(file_path, output_dir):
    """
    Converte um arquivo PDF ou PPTX em imagens PNG usando a abordagem mais robusta.
    - Para PDF: Usa PyMuPDF para renderização direta, sem dependências externas.
    - Para PPTX: Mantém o uso de LibreOffice para converter para PDF primeiro.
    Retorna uma lista de caminhos completos das imagens geradas.
    """
    try:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        ext = os.path.splitext(file_path)[1].lower()
        image_paths = []

        pdf_to_process = None

        if ext == '.pptx':
            # A conversão de PPTX para PDF ainda depende do LibreOffice
            pdf_file_path = os.path.join(output_dir, base_name + ".pdf")
            try:
                # LibreOffice's launch script needs standard Unix tools (dirname, basename, etc.)
                # When running under systemd, PATH may be too restricted, so we ensure it's complete
                env = os.environ.copy()
                env['PATH'] = '/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin:' + env.get('PATH', '')
                subprocess.run(
                    ["/usr/bin/libreoffice", "--headless", "--convert-to", "pdf", file_path, "--outdir", output_dir],
                    check=True, timeout=120, env=env
                )
                if not os.path.exists(pdf_file_path):
                    raise FileNotFoundError(f"A conversão de PPTX para PDF falhou em gerar o arquivo: {pdf_file_path}")
                pdf_to_process = pdf_file_path
            except subprocess.CalledProcessError as e:
                raise RuntimeError(f"Erro ao converter PPTX para PDF com LibreOffice: {e}")
            except subprocess.TimeoutExpired:
                raise RuntimeError("A conversão de PPTX para PDF demorou muito (timeout).")

        elif ext == '.pdf':
            pdf_to_process = file_path
        
        else:
            raise ValueError(f"Formato de arquivo não suportado para conversão em imagem: {ext}")

        # Processamento do PDF usando PyMuPDF (muito mais robusto)
        if pdf_to_process:
            doc = fitz.open(pdf_to_process)
            for i, page in enumerate(doc):
                # Renderiza a página como uma imagem (pixmap)
                # Aumentar o DPI melhora a qualidade do OCR
                pix = page.get_pixmap(dpi=300) 
                image_path = os.path.join(output_dir, f"{base_name}-page-{i}.png")
                pix.save(image_path)
                image_paths.append(image_path)
            doc.close()

        # Se a origem era um PPTX, remove o PDF intermediário
        if ext == '.pptx' and pdf_to_process and os.path.exists(pdf_to_process):
            os.remove(pdf_to_process)

        if not image_paths:
            raise RuntimeError("Nenhuma imagem foi gerada a partir do arquivo de entrada.")
            
        return image_paths

    except Exception as e:
        # Log do erro para depuração
        logger.error(f"Falha crítica na conversão para imagem: {e}", exc_info=True)
        raise RuntimeError(f"Erro ao converter arquivo para imagens: {e}")

def generate_thumbnail(pdf_path: str, document_id: int) -> str:
    """
    Gera uma thumbnail da primeira página de um PDF e a salva.
    Retorna o caminho relativo da thumbnail salva.
    """
    try:
        doc = fitz.open(pdf_path)
        page = doc.load_page(0)  # Carrega a primeira página
        pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5))  # Reduz a resolução para thumbnail
        
        # Garante que o diretório de thumbnails exista
        os.makedirs(Config.THUMBNAIL_FOLDER, exist_ok=True)
        
        thumbnail_filename = f"thumb_{document_id}_{os.path.basename(pdf_path)}.png"
        thumbnail_path = os.path.join(Config.THUMBNAIL_FOLDER, thumbnail_filename)
        
        pix.save(thumbnail_path)
        doc.close()
        
        logger.info(f"Thumbnail gerada: {thumbnail_path}")
        # Retorna o caminho relativo para ser salvo no banco de dados e usado pelo frontend
        return os.path.join(Config.STATIC_URL_PATH_PREFIX.strip('/'), "uploads", "thumbnails", thumbnail_filename)
    except Exception as e:
        logger.error(f"Erro ao gerar thumbnail para {pdf_path}: {e}", exc_info=True)
        return None

def aplicar_ocr_em_todas_imagens(image_paths):
    """
    Aplica OCR a cada imagem na lista e retorna o texto concatenado de todas as páginas.
    """
    try:
        textos = []
        for image_path in image_paths:
            text = pytesseract.image_to_string(Image.open(image_path), lang='por+spa+eng')
            textos.append(text)
        return "\n".join(textos)  # Concatena os textos com quebras de linha
    except Exception as e:
        raise RuntimeError(f"Erro ao aplicar OCR nas imagens: {e}")

# Função para obter dados de entrada (movida de routes.py)
def get_input_data(filepath, ext):
    try:
        if ext in ['mp3', 'wav', 'm4a', 'mp4', 'avi', 'mov']:
            logger.debug(f"Transcrevendo: {filepath}")
            return transcribe_audio(filepath)
        elif ext in ['pdf', 'pptx']:
            logger.debug(f"Convertendo {ext} para imagens: {filepath}")
            image_paths = converter_para_imagem(filepath, Config.UPLOAD_FOLDER)
            return aplicar_ocr_em_todas_imagens(image_paths)
        elif ext == 'docx':
            logger.debug(f"Extraindo texto de DOCX: {filepath}")
            import docx
            doc = docx.Document(filepath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        logger.debug(f"Lendo texto: {filepath}")
        with open(filepath, 'r', encoding='utf-8') as f: # Adicionado encoding
             return f.read()
    except Exception as e:
        logger.error(f"Erro ao obter dados: {str(e)}", exc_info=True)
        raise


def extract_text_direct(filepath: str) -> str:
    """
    Extract text directly from PDF using PyMuPDF (no OCR, no temp files).
    Instant for digital PDFs. Returns empty string for scanned PDFs.
    """
    try:
        doc = fitz.open(filepath)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"Direct text extraction failed for {filepath}: {e}")
        return ""


def extract_pages_from_pdf(filepath: str, ocr_threshold: int = 100):
    """
    Extrai texto PÁGINA A PÁGINA de um PDF. Retorna [(page_number, text), ...].

    Para cada página tenta o texto nativo (PyMuPDF, instantâneo). Se a página tiver menos
    que `ocr_threshold` chars (página escaneada/imagem), faz OCR só daquela página — então
    PDF MISTO (digital + escaneado) é tratado corretamente, ao contrário do antigo limiar
    global tudo-ou-nada. O número de página vai pra metadata do chunk (habilita citação).
    """
    pages = []
    tmp_imgs = []
    try:
        doc = fitz.open(filepath)
        for i, page in enumerate(doc):
            text = (page.get_text() or "").strip()
            if len(text) < ocr_threshold:
                try:
                    pix = page.get_pixmap(dpi=300)  # DPI alto melhora o OCR
                    os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
                    img_path = os.path.join(
                        Config.UPLOAD_FOLDER, f"_ocr_{os.getpid()}_p{i}.png"
                    )
                    pix.save(img_path)
                    tmp_imgs.append(img_path)
                    ocr_text = pytesseract.image_to_string(
                        Image.open(img_path), lang='por+spa+eng'
                    ).strip()
                    if len(ocr_text) > len(text):
                        text = ocr_text
                except Exception as e_ocr:
                    logger.warning(f"OCR da página {i + 1} de {filepath} falhou: {e_ocr}")
            pages.append((i + 1, text))
        doc.close()
        return pages
    except Exception as e:
        logger.error(f"Falha na extração por página de {filepath}: {e}", exc_info=True)
        # Fallback: texto direto inteiro como página única.
        return [(1, extract_text_direct(filepath))]
    finally:
        limpar_arquivos_locais(tmp_imgs)


def extract_pages_from_pptx(filepath: str):
    """
    Extrai texto por SLIDE de um PPTX (cada slide vira uma 'página'). Converte para imagens
    (LibreOffice→PDF→PyMuPDF) e aplica OCR em cada slide. Retorna [(slide_number, text), ...].
    """
    image_paths = []
    try:
        image_paths = converter_para_imagem(filepath, Config.UPLOAD_FOLDER)
        pages = []
        for i, img_path in enumerate(image_paths):
            try:
                text = pytesseract.image_to_string(
                    Image.open(img_path), lang='por+spa+eng'
                ).strip()
            except Exception as e_ocr:
                logger.warning(f"OCR do slide {i + 1} de {filepath} falhou: {e_ocr}")
                text = ""
            pages.append((i + 1, text))
        return pages
    except Exception as e:
        logger.error(f"Falha na extração por slide de {filepath}: {e}", exc_info=True)
        return []
    finally:
        limpar_arquivos_locais(image_paths)


def _is_likely_medical_image(image_bytes: bytes, width: int, height: int) -> bool:
    """
    Heuristic filter to discard non-medical images BEFORE sending to vision API.

    Rejects: solid colors, gradients, near-blank backgrounds, page-sized captures,
    logos, license icons, decorative banners, small graphics with flat fills.

    Accepts: clinical photos, radiographs, histology, ECGs, growth charts,
    anatomical diagrams — anything with real visual complexity.
    """
    import io
    import numpy as np

    # Hard minimum: file size < 5KB is almost certainly an icon/logo
    if len(image_bytes) < 5_000:
        return False

    # Reject very narrow/tall or very wide/short images (banners, logos, bars)
    aspect = max(width, height) / max(min(width, height), 1)
    if aspect > 4.0:
        return False

    # Minimum meaningful size for medical content: at least 250px on shorter side
    if min(width, height) < 250:
        return False

    try:
        img = Image.open(io.BytesIO(image_bytes))

        # Convert to RGB if needed (CMYK PDFs, palette modes, etc.)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')

        # Reject page-sized images (full-page backgrounds/scans)
        # A4 at 150dpi ~= 1240x1754; at 300dpi ~= 2480x3508
        if width * height > 2_000_000 and 1.2 < aspect < 1.6:
            small = img.resize((64, 64))
            arr = np.array(small)
            if arr.mean() > 235:  # Nearly all white — blank page
                return False

        # Downsample for fast analysis
        small = img.resize((64, 64))
        arr = np.array(small, dtype=np.float32)

        # Grayscale images
        if len(arr.shape) == 2:
            std = arr.std()
            if std < 15:
                return False
            return True

        # Color images — check per-channel std
        channel_stds = arr.std(axis=(0, 1))
        mean_std = channel_stds.mean()

        # Very low variance = solid fill, gradient, blank
        if mean_std < 15:
            return False

        # Check color diversity via quantized unique colors
        pixels = arr.reshape(-1, 3).astype(np.uint8)
        quantized = (pixels // 32) * 32
        unique_colors = len(set(map(tuple, quantized.tolist())))

        # Fewer than 15 unique colors = flat graphic (logo, icon, simple diagram)
        if unique_colors < 15:
            return False

        return True

    except Exception as e:
        logger.debug(f"[IMAGES] Heuristic filter error: {e}")
        return True


def extract_images_from_pdf(pdf_path: str, document_id: int, library_id: int, output_dir: str) -> list:
    """
    Extract embedded images from a PDF using PyMuPDF.
    Filters out small images AND non-medical content (solid fills, gradients,
    backgrounds, page captures) using heuristic analysis.
    Returns list of dicts with image metadata for DB insertion.

    This runs in a thread pool (blocking I/O).
    """
    import uuid as _uuid

    results = []
    skipped_heuristic = 0
    os.makedirs(output_dir, exist_ok=True)

    try:
        doc = fitz.open(pdf_path)
        image_count = 0

        for page_num in range(len(doc)):
            if image_count >= Config.MAX_IMAGES_PER_DOCUMENT:
                logger.info(f"[IMAGES] Document {document_id}: reached max images limit ({Config.MAX_IMAGES_PER_DOCUMENT})")
                break

            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                if image_count >= Config.MAX_IMAGES_PER_DOCUMENT:
                    break

                xref = img_info[0]
                try:
                    extracted = doc.extract_image(xref)
                    if not extracted:
                        continue

                    width = extracted.get("width", 0)
                    height = extracted.get("height", 0)
                    image_bytes = extracted.get("image")
                    ext_type = extracted.get("ext", "png")

                    if not image_bytes:
                        continue

                    # Filter 1: dimension/area check
                    if width < Config.MIN_IMAGE_DIMENSION or height < Config.MIN_IMAGE_DIMENSION:
                        continue
                    if width * height < Config.MIN_IMAGE_AREA:
                        continue

                    # Filter 2: heuristic — reject solid fills, gradients, blank pages
                    if not _is_likely_medical_image(image_bytes, width, height):
                        skipped_heuristic += 1
                        continue

                    # Generate unique filename
                    short_uuid = _uuid.uuid4().hex[:8]
                    filename = f"{document_id}_p{page_num}_i{img_index}_{short_uuid}.{ext_type}"
                    filepath = os.path.join(output_dir, filename)

                    with open(filepath, "wb") as f:
                        f.write(image_bytes)

                    file_size = len(image_bytes)

                    results.append({
                        "document_id": document_id,
                        "library_id": library_id,
                        "image_filename": filename,
                        "page_number": page_num,
                        "image_index": img_index,
                        "width": width,
                        "height": height,
                        "file_size_bytes": file_size,
                    })
                    image_count += 1

                except Exception as e:
                    logger.warning(f"[IMAGES] Failed to extract image xref={xref} from page {page_num} of doc {document_id}: {e}")
                    continue

        doc.close()
        logger.info(f"[IMAGES] Document {document_id}: extracted {len(results)} images, skipped {skipped_heuristic} by heuristic filter")

    except Exception as e:
        logger.error(f"[IMAGES] Failed to extract images from document {document_id}: {e}", exc_info=True)

    return results


def get_input_data_from_bytes(content_bytes: bytes, ext: str) -> str:
    """Extract text content from file bytes. Used by copilot for inline document analysis."""
    try:
        if ext == 'pdf':
            # Extract text directly from PDF bytes using PyMuPDF
            doc = fitz.open(stream=content_bytes, filetype="pdf")
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            text = "\n".join(text_parts).strip()
            # If PDF has extractable text, return it
            if len(text) > 50:
                return text
            # Fallback: scanned PDF → save to temp, use OCR flow
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            try:
                return get_input_data(tmp_path, ext)
            finally:
                os.unlink(tmp_path)
        elif ext == 'docx':
            # Save to temp and extract with python-docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            try:
                import docx
                doc = docx.Document(tmp_path)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            finally:
                os.unlink(tmp_path)
        elif ext in ['txt', 'md', 'csv', 'json', 'xml', 'html']:
            return content_bytes.decode('utf-8')
        else:
            # Generic: try UTF-8 decode
            return content_bytes.decode('utf-8')
    except Exception as e:
        logger.error(f"Erro ao extrair texto de bytes ({ext}): {str(e)}", exc_info=True)
        raise
